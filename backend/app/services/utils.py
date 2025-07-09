import os
import re
import pickle
import spacy
import nltk
from nltk.corpus import stopwords
from PyPDF2 import PdfReader
from docx import Document
import io
from typing import List


NLTK_DATA_PATH = os.path.join(
    os.path.dirname(__file__),
    "../model/nltk_data",
)
if not os.path.exists(
    NLTK_DATA_PATH,
):
    os.makedirs(
        NLTK_DATA_PATH,
    )

nltk.data.path.append(
    NLTK_DATA_PATH,
)
nltk.download(
    "punkt",
    download_dir=NLTK_DATA_PATH,
)
nltk.download(
    "stopwords",
    download_dir=NLTK_DATA_PATH,
)

stop_words = set(
    stopwords.words(
        "english",
    )
)
nlp = spacy.load(
    "en_core_web_sm",
)

clf = pickle.load(
    open(
        os.path.join(
            os.path.dirname(__file__),
            "../model/best_model.pkl",
        ),
        "rb",
    )
)
tfidf_vectorizer = pickle.load(
    open(
        os.path.join(
            os.path.dirname(__file__),
            "../model/tfidf.pkl",
        ),
        "rb",
    )
)


def process_document(file_bytes, file_name):
    file_extension = os.path.splitext(file_name)[1].lower()
    raw_text = ""
    try:
        if file_extension == ".txt" or file_extension == ".md":
            raw_text = file_bytes.decode()

        elif file_extension == ".pdf":
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                raw_text += page.extract_text() or ""

        elif file_extension == ".docx":
            doc = Document(io.BytesIO(file_bytes))

            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    raw_text += para.text + "\n"

            for para in doc.paragraphs:
                raw_text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace("\n", " ")
                        row_text.append(cell_text)
                    raw_text += "\t".join(row_text) + "\n"

            for section in doc.sections:
                footer = section.footer
                for para in footer.paragraphs:
                    raw_text += para.text + "\n"

        else:
            print(
                f"Unsupported file type: {file_extension}. Please upload TXT, MD, PDF, or DOCX."
            )
            return None

    except Exception as e:
        print(f"Error processing file {file_name}: {e}")
        return None

    return raw_text


def clean_resume(txt):
    cleantxt = re.sub(r"https\\S+", "", txt)
    cleantxt = re.sub(r"@\\S+|#\\S+", "", cleantxt)
    cleantxt = re.sub(r"[^\w\s]", "", cleantxt)
    doc = nlp(cleantxt)
    tokens = [
        token.lemma_.lower() for token in doc if token.text.lower() not in stop_words
    ]
    return " ".join(tokens)


def extract_name_and_email(text):
    email_regex = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    emails = re.findall(email_regex, text)
    lines = text.splitlines()
    name = lines[0].strip() if lines else "N/A"
    email = emails[0] if emails else "N/A"
    return name, email


def extract_contact_number_from_resume(text):
    contact_number = None
    pattern = r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"
    match = re.search(pattern, text)
    if match:
        contact_number = match.group()
    return contact_number


def extract_college_name(text):
    lines = text.split("\n")
    college_pattern = r"(?i).*(college|university|institute).*"
    for line in lines:
        if re.match(college_pattern, line):
            college_name = line.strip()
            max_length = 128
            return college_name[:max_length]
    return ""


def extract_education_info(text):
    pattern = r"(?i)(?:(?:Bachelor|B\.S\.|B\.A\.|Master|M\.S\.|M\.A\.|Ph\.D\.)\s(?:[A-Za-z]+\s)*[A-Za-z]+)"
    matches = re.findall(pattern, text)
    education = [match.strip() for match in matches]
    education_text = ", ".join(education) if education else " "
    return education_text


def extract_work_experience(text: str) -> List[str]:
    work_experience_keywords = [
        r"\b(intern|manager|developer|engineer|analyst|lead|consultant|specialist|director|officer|administrator|associate|scientist|technician|programmer|designer|researcher|trainee|staff)\b",
        r"\b(Senior|Junior|Lead|Principal|Entry|Mid-level|Experienced)\s+\w+\b",
    ]
    pattern = re.compile("|".join(work_experience_keywords), re.IGNORECASE)
    work_experience_info = []
    for line in text.splitlines():
        if pattern.search(line):
            work_experience_info.append(line.strip())
    return work_experience_info if work_experience_info else []


def extract_projects(text: str) -> List[str]:
    projects_info = []
    in_project_section = False
    project_section_keywords = [
        "PROJECTS",
        "PERSONAL PROJECTS",
        "ACADEMIC PROJECTS",
        "PROJECT EXPERIENCE",
    ]
    section_end_keywords = [
        "EXPERIENCE",
        "EDUCATION",
        "SKILLS",
        "CERTIFICATIONS",
        "AWARDS",
    ]
    lines = text.splitlines()
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        is_project_header = any(
            re.search(r"\b" + keyword + r"\b", stripped_line, re.IGNORECASE)
            for keyword in project_section_keywords
        )

        is_section_ender = any(
            re.search(r"\b" + keyword + r"\b", stripped_line, re.IGNORECASE)
            for keyword in section_end_keywords
        )

        if is_project_header:
            in_project_section = True
            continue

        if in_project_section:
            if is_section_ender and not is_project_header:
                in_project_section = False
                break
            projects_info.append(stripped_line)

    if not projects_info:
        project_line_keywords = [
            "developed",
            "implemented",
            "created",
            "designed",
            "project on",
            "built a",
            "application for",
        ]
        for line in lines:
            if any(keyword in line.lower() for keyword in project_line_keywords):
                if len(line.strip()) > 20:
                    projects_info.append(line.strip())
                    if len(projects_info) >= 10:
                        break

    return projects_info


def extract_skills_from_resume(text, skills_list):
    skills = []
    for skill in skills_list:
        pattern = r"\b{}\b".format(re.escape(skill))
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            skills.append(skill)
    return skills


def predict_category(cleaned_resume):
    input_features = tfidf_vectorizer.transform([cleaned_resume])
    prediction_id = clf.predict(input_features)[0]
    category_mapping = {
        15: "Java Developer",
        23: "Testing",
        8: "DevOps Engineer",
        20: "Python Developer",
        24: "Web Designing",
        12: "HR",
        13: "Hadoop",
        3: "Blockchain",
        10: "ETL Developer",
        18: "Operations Manager",
        6: "Data Science",
        22: "Sales",
        16: "Mechanical Engineer",
        1: "Arts",
        7: "Database",
        11: "Electrical Engineering",
        14: "Health and fitness",
        19: "PMO",
        4: "Business Analyst",
        9: "DotNet Developer",
        2: "Automation Testing",
        17: "Network Security Engineer",
        21: "SAP Developer",
        5: "Civil Engineer",
        0: "Advocate",
    }
    return category_mapping.get(prediction_id, "Unknown")


def is_valid_resume(text):
    if not text:
        return False
    resume_keywords = [
        "Experience",
        "Education",
        "Skills",
        "Profile",
        "Work History",
        "Projects",
        "Certifications",
    ]

    if any(re.search(keyword, text, re.I) for keyword in resume_keywords):
        return True

    return False
